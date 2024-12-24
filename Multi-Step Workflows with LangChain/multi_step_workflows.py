import os
import re
import nltk
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from fpdf import FPDF
from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
from nltk.tokenize import sent_tokenize

nltk.download("punkt")

# load the summarization model
def initialize_summarizer():
    model_name = "facebook/bart-large-cnn"
    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
    return pipeline("summarization", model=model, tokenizer=tokenizer)

# extract text from a document
def load_document(file_path):
    content = ""
    if file_path.endswith(".pdf"):
        reader = PdfReader(file_path)
        for page in reader.pages:
            content += page.extract_text() if page.extract_text() else ""
    elif file_path.endswith(".docx"):
        doc = DocxDocument(file_path)
        for para in doc.paragraphs:
            content += para.text
    else:
        raise ValueError("invalid document!")
    return content

# text preprocessing
def preprocess_text(text):
    text = text.strip()
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"[^\w\s.,!?]", "", text)
    text = re.sub(r"\n+", " ", text)
    return text

# split the text to paragraphs
def split_into_paragraphs(text, paragraph_size=3):
    sentences = sent_tokenize(text)
    paragraphs = [" ".join(sentences[i:i + paragraph_size]) for i in range(0, len(sentences), paragraph_size)]
    return paragraphs

# summarizing paragraphs
def summarize_paragraphs(summarizer, paragraphs, max_length=70, min_length=35):
    summaries = []
    for paragraph in paragraphs:
        if len(paragraph.strip()) > 0:
            summary = summarizer(paragraph, max_length=max_length, min_length=min_length, do_sample=False)
            summaries.append(summary[0]['summary_text'])
    return " ".join(summaries)

def save_summary_to_word(summary_text, output_path):
    doc = DocxDocument()
    doc.add_heading('Summary', 0)
    doc.add_paragraph(summary_text)
    doc.save(output_path)

def save_summary_to_pdf(summary_text, output_path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, summary_text)
    pdf.output(output_path)

if __name__ == "__main__":
    input_file = "input file path"
    output_word_file = "output path"
    output_pdf_file = "output path"

    raw_text = load_document(input_file)
    if not raw_text.strip():
        raise ValueError("invalid content")

    clean_text = preprocess_text(raw_text)

    paragraphs = split_into_paragraphs(clean_text, paragraph_size=5)

    summarizer = initialize_summarizer()
    final_summary = summarize_paragraphs(summarizer, paragraphs)

    save_summary_to_word(final_summary, output_word_file)
    save_summary_to_pdf(final_summary, output_pdf_file)
    print("Done!")